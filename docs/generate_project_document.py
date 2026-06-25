from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Copic_Project_Explained.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C3D0"
MUTED = "555555"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.type = "dxa"
            tc_w.w = int(width * 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def style_table(table, widths: list[float], header_fill: str = LIGHT_BLUE) -> None:
    set_table_width(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, widths, header_fill)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(f" {body}")
    set_table_width(table, [6.5])
    set_cell_shading(cell, LIGHT_GRAY)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Copic project explained"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)

    footer = section.footer.paragraphs[0]
    footer.text = "Generated from the local project workspace"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


PAGES = [
    ("/", "Landing page", "Public product entry with worker/client positioning and calls to action."),
    ("/auth/register", "Register", "Email/password registration with role selection."),
    ("/auth/login", "Login", "Public sign-in page, redirects signed-in users to role home."),
    ("/auth/admin", "Admin login", "Hidden admin username/password login flow."),
    ("/verify-email", "Email verification", "6-digit Resend OTP verification screen."),
    ("/complete-profile", "Profile completion", "Fallback profile creation/completion path after auth."),
    ("/dashboard", "Worker dashboard", "Worker overview, skills, ratings, wallet summary, applications/live/completed job modal."),
    ("/jobs", "Jobs marketplace", "Worker-facing open jobs list with search, location and pay sorting."),
    ("/jobs/[id]", "Job details", "Worker application form and job detail view."),
    ("/jobs/new", "New job", "Client job-posting wizard."),
    ("/find-work", "Client work management", "Client posted work dashboard, rehire controls and recurring job status."),
    ("/applications", "Applications", "Client applicant review and worker application/current-job workflow."),
    ("/workers", "Worker directory", "Client search/discovery page for worker profiles."),
    ("/chat", "Chat", "Accepted job conversations and message threads."),
    ("/wallet", "Wallet", "Balances, deposits, withdrawals, transactions and payment setup prompts."),
    ("/payment-setup", "Payment setup", "Dedicated wallet/payment readiness screen."),
    ("/notifications", "Notifications", "Alerts, archive/delete actions and rating prompts."),
    ("/profile", "Profile", "Public-ish account profile editing and display."),
    ("/settings", "Verification settings", "KYC document, address and identity submission flow."),
    ("/account-settings", "Account settings", "Email and phone account edits."),
    ("/help", "Help", "Support ticket creation and user support conversation view."),
    ("/faq", "FAQ", "Published FAQ records for users."),
    ("/about", "About", "Public company/product description."),
    ("/offline", "Offline fallback", "PWA offline page."),
    ("/admin", "Admin dashboard", "Admin overview and control hub."),
    ("/admin/users", "Admin users", "User lock/unlock and moderation controls."),
    ("/admin/kyc", "Admin KYC", "Verification review queue."),
    ("/admin/transactions", "Admin transactions", "Finance and M-Pesa reconciliation surface."),
    ("/admin/support", "Admin support", "Support ticket response and status management."),
    ("/admin/audit", "Admin audit", "Security/admin action logs."),
    ("/admin/disputes", "Admin disputes", "Dispute management surface."),
    ("/admin/reports", "Admin reports", "Platform reporting page."),
]


APIS = [
    ("GET/PATCH/DELETE", "/api/jobs", "Loads, edits, completes, cancels or deletes jobs. Blocks deletion after acceptance."),
    ("POST", "/api/jobs/create", "Creates posted work, validates job inputs, reserves escrow where required."),
    ("POST", "/api/jobs/rehire", "Creates direct rehire jobs and pending rehire applications with explicit start dates."),
    ("GET/POST/PATCH", "/api/applications", "Lists applications, creates applications, accepts/cancels/completes/rehire-responds."),
    ("GET/POST", "/api/chat", "Lists conversations/messages and sends messages into unlocked conversations."),
    ("GET/PATCH", "/api/notifications", "Lists, archives, restores, and deletes user notifications."),
    ("GET/POST", "/api/ratings", "Loads aggregate reviews and creates ratings after completed jobs."),
    ("GET", "/api/users", "Lists worker/user records for discovery surfaces."),
    ("GET", "/api/wallet", "Loads SQL-mode wallet and transactions."),
    ("POST", "/api/wallet/test-deposit", "Development-only test deposit when enabled."),
    ("POST", "/api/wallet/withdraw", "Creates withdrawal requests and validates wallet balance."),
    ("POST", "/api/mpesa/stk-push", "Starts Daraja STK Push deposits and creates pending transactions."),
    ("POST", "/api/mpesa/callback", "Receives M-Pesa callback and credits wallet after validation."),
    ("POST", "/api/kyc/start", "Creates Didit hosted KYC sessions and verification records."),
    ("POST", "/api/kyc/didit-webhook", "Validates Didit webhook signatures and updates identity status."),
    ("GET/POST", "/api/auth/me", "Loads current profile and creates/imports SQL-mode profiles."),
    ("POST", "/api/auth/create-profile", "Creates Firestore user profile after registration."),
    ("POST", "/api/auth/send-email-otp", "Creates and emails a short-lived OTP."),
    ("POST", "/api/auth/verify-email-otp", "Verifies OTP hash, marks auth/profile email verified."),
    ("POST", "/api/auth/admin-login", "Validates hidden admin credentials and profile."),
    ("PATCH", "/api/account-settings", "Updates account email and phone details."),
    ("GET/POST/DELETE", "/api/profile/skills", "Loads, creates/updates, and deletes worker skill profiles."),
    ("GET", "/api/activities", "Loads user activity feed."),
    ("GET/POST/PATCH", "/api/admin/tickets and /[id]", "Admin support ticket list, replies and status changes."),
    ("GET/POST", "/api/admin/users/actions/audit/stats", "Admin users, actions, audit logs and dashboard statistics."),
    ("POST", "/api/admin/mpesa/reconcile", "Manual finance reconciliation for stuck M-Pesa transactions."),
]


SERVICES = [
    ("services/auth.ts", "Client auth wrappers: profile creation, email login/register, admin login, phone linking, logout."),
    ("services/jobs.ts", "Client job/application API layer: subscribe jobs, apply, accept, cancel, complete, rehire."),
    ("services/wallet.ts", "Wallet subscriptions plus deposit, withdrawal, test deposit and legacy callable wrappers."),
    ("services/chat.ts", "Conversation/message polling and send-message API calls."),
    ("services/notifications.ts", "Push setup, foreground messages, notification list and archive/delete operations."),
    ("services/kyc.ts", "Uploads verification files and starts/reviews verification."),
    ("services/worker-skills.ts", "Loads, saves and deletes detailed worker skill profiles."),
    ("services/ratings.ts", "Loads rating history and submits user/client ratings."),
    ("services/support.ts", "Support tickets, messages, FAQ publishing and subscriptions."),
    ("services/users.ts", "Worker directory data loading and normalization."),
    ("services/activities.ts", "User activity feed loading."),
    ("services/emailVerification.ts", "Email-verification guard message/helper."),
]


LIBS = [
    ("lib/firebase.ts", "Firebase client initialization and require* helpers for auth, Firestore, Storage, Functions and Messaging."),
    ("lib/firebase-admin.ts / firebaseAdmin.ts", "Firebase Admin app, Auth and Firestore helpers for server routes."),
    ("lib/server-auth.ts", "Server-side Firebase token verification and email-verification enforcement."),
    ("lib/admin-security.ts", "Admin permission model, admin actor lookup and audit logging."),
    ("lib/admin-credentials.ts", "Local admin password hashing and stored admin credential helpers."),
    ("lib/data-backend.ts", "Switches Firebase/SQL behavior through DATA_BACKEND/NEXT_PUBLIC_DATA_BACKEND."),
    ("lib/local-sql.ts", "SQLite schema, migrations, local users/jobs/applications/wallet/chat/notifications/ratings/admin settings."),
    ("lib/mpesa.ts", "M-Pesa phone normalization, account references, access tokens, STK Push and status query helpers."),
    ("lib/didit.ts", "Didit session creation, webhook signature verification and national ID hashing."),
    ("lib/email-otp.ts", "Email normalization, OTP validation, hashing and authenticated OTP user checks."),
    ("lib/location.ts", "Default Kenya location object."),
    ("lib/jobCategories.ts", "Category groups and flattened category options."),
    ("lib/server-debug.ts", "Server-side debug logging wrapper."),
]


UTILS = [
    ("utils/validation.ts", "Zod schemas for job and profile inputs."),
    ("utils/duration.ts", "Duration units, conversion to hours, labels and job duration extraction."),
    ("utils/pricing.ts", "Worker-visible pay calculation after platform fee."),
    ("utils/money.ts", "Platform fee rate, service fee, worker net and KES formatting."),
    ("utils/phone.ts", "Kenyan phone normalization."),
    ("utils/jobUnits.ts", "Quantity/unit display and payment method options."),
    ("utils/jobCategoryMatcher.ts", "Maps job titles to likely categories."),
    ("utils/verification.ts", "Verification status normalization and labels."),
]


COMPONENTS = [
    ("components/layout/Shell.tsx", "Role-aware app shell, navigation and layout chrome."),
    ("components/auth/AuthForm.tsx", "Shared login/register form UI."),
    ("components/auth/EmailVerificationRequired.tsx", "Email verification guard UI."),
    ("components/jobs/PostWorkWizard.tsx", "Client job creation workflow."),
    ("components/jobs/JobCard.tsx", "Reusable job card for marketplace and client views."),
    ("components/location/MapPicker.tsx", "Mapbox address and coordinate selector."),
    ("components/profile/AddSkillModal.tsx", "Worker skill add/edit modal with proof and pricing details."),
    ("components/ratings/RatingHistory.tsx", "Ratings and reviews display."),
    ("components/wallet/PaymentSetupCard.tsx", "Payment readiness card."),
    ("components/pwa/PwaBootstrap.tsx", "Service worker and install/bootstrap behavior."),
    ("components/ui/*", "Shared buttons, cards, modals, loading, skeleton, empty states and star rating input."),
]


MODELS = [
    ("Role/AdminRole/AdminPermission", "Worker, client and admin role variants plus granular admin permission names."),
    ("UserProfile", "Main account record with role, display name, skills, location, ratings, verification, lock and wallet metadata."),
    ("WorkerSkillProfile", "Detailed worker skill record with category, level, proof type, charge model and rating stats."),
    ("LocationFields", "County, town, estate/area, landmark, address text and coordinates."),
    ("VerificationRecord", "Identity/address verification submission, uploaded document URLs, provider status and review result."),
    ("Job", "Posted work record with client, pay, duration, quantity, skills, recurrence, rehire and status fields."),
    ("Application", "Worker-to-job relationship with status, enriched worker/client display data and cover note."),
    ("Conversation/Message", "Accepted-job chat container and message items."),
    ("Wallet/Transaction", "Available/pending balances, deposits, withdrawals, payouts, wallet credits and fees."),
    ("Rating", "Post-job review from one user to another."),
    ("AppNotification/Activity", "User alerts and activity feed items."),
    ("AdminAuditLog/WalletAdjustment", "Admin security history and finance adjustments."),
]


ENV_ROWS = [
    ("Firebase public", "NEXT_PUBLIC_FIREBASE_*", "Client Firebase app, auth domain, project, storage, messaging and app ID."),
    ("Firebase admin", "FIREBASE_PROJECT_ID, FIREBASE_CLIENT_EMAIL, FIREBASE_PRIVATE_KEY", "Server-side Admin SDK credentials."),
    ("Data mode", "NEXT_PUBLIC_DATA_BACKEND, DATA_BACKEND, LOCAL_SQLITE_PATH", "Enables local SQLite development mode."),
    ("Admin bootstrap", "ADMIN_USERNAME, ADMIN_PASSWORD, ADMIN_EMAIL", "Hidden admin login defaults/settings."),
    ("Email OTP", "RESEND_API_KEY, RESEND_FROM_EMAIL, OTP_SECRET", "Email verification code delivery and hash protection."),
    ("M-Pesa", "MPESA_CONSUMER_KEY, MPESA_CONSUMER_SECRET, MPESA_SHORTCODE, MPESA_PASSKEY, MPESA_CALLBACK_URL, MPESA_ENV", "Daraja STK Push and callback processing."),
    ("KYC", "KYC_PROVIDER, DIDIT_API_KEY, DIDIT_WORKFLOW_ID, DIDIT_WEBHOOK_SECRET", "Hosted identity verification."),
    ("Maps", "NEXT_PUBLIC_MAPBOX_TOKEN / NEXT_PUBLIC_GOOGLE_MAPS_API_KEY", "Address and coordinate picking."),
    ("Platform settings", "NEXT_PUBLIC_PLATFORM_FEE_RATE, NEXT_PUBLIC_MPESA_PAYBILL, ALLOW_TEST_DEPOSITS", "Fee display, payment identity and development finance controls."),
]


def file_inventory() -> list[list[str]]:
    rows: list[list[str]] = []
    skip_suffixes = {".map", ".tsbuildinfo", ".lock"}
    for path in sorted(ROOT.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT).as_posix()
        if "node_modules" in rel or ".next" in rel or rel.startswith("docs/"):
            continue
        if any(rel.endswith(s) for s in skip_suffixes):
            continue
        if rel in {"package-lock.json"}:
            continue
        if rel.startswith("functions/lib/"):
            continue
        if rel.startswith("public/design/") or rel.startswith("public/icons/"):
            purpose = "Static design/icon asset."
        elif rel.startswith("app/api/"):
            purpose = "Server route handler."
        elif rel.startswith("app/"):
            purpose = "App Router page, layout, manifest or global style."
        elif rel.startswith("components/"):
            purpose = "Reusable UI/component module."
        elif rel.startswith("services/"):
            purpose = "Client-side data/API service."
        elif rel.startswith("lib/"):
            purpose = "Shared infrastructure/helper library."
        elif rel.startswith("utils/"):
            purpose = "Pure utility or validation helper."
        elif rel.startswith("functions/"):
            purpose = "Firebase Cloud Functions source/config."
        elif rel.startswith("public/"):
            purpose = "Public static/PWA asset."
        elif rel.endswith(".rules"):
            purpose = "Firebase security rules."
        else:
            purpose = "Project configuration or documentation."
        rows.append([rel, purpose])
    return rows


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Copic Project Explained")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("A practical handover document for the Temp gig-work marketplace app")
    meta = doc.add_paragraph()
    meta.add_run("Generated from: ").bold = True
    meta.add_run(str(ROOT))
    meta.add_run("\nDocument purpose: ").bold = True
    meta.add_run("Explain every major part of the codebase, user workflow, data model, API surface, security layer and deployment setup.")

    doc.add_heading("Quick Orientation", level=1)
    doc.add_paragraph(
        "Copic is a mobile-first Progressive Web App for temporary work. Clients post jobs and hire workers; "
        "workers apply, accept direct hires, complete work, and receive wallet payments. The app uses Next.js App Router "
        "for UI and API routes, Firebase for production identity/data/storage/hosting/functions, and SQLite for local development."
    )
    add_note(
        doc,
        "Core idea:",
        "The browser handles role-specific screens and forms, Next.js route handlers enforce server rules, and Firebase or local SQLite stores the operational records."
    )
    add_table(doc, ["Area", "What it does"], [
        ["Workers", "Find jobs, apply, cancel pending applications, receive direct hire requests, chat after acceptance, mark jobs complete, view ratings and wallet status."],
        ["Clients", "Post work, review applicants, accept workers, escrow funds, confirm completion, rehire workers, manage support and payments."],
        ["Admins", "Review users/KYC/support/transactions, manage moderation, audit security-sensitive actions and reconcile finance records."],
        ["Platform", "Enforces roles, email/KYC/payment state, wallet accounting, service fees, notification delivery, and deployment security rules."],
    ], [1.5, 5.0])

    doc.add_heading("Technology Stack", level=1)
    add_table(doc, ["Layer", "Implementation"], [
        ["Frontend", "Next.js 16 App Router, React 19, TypeScript, Tailwind CSS, lucide-react icons, Framer Motion, Sonner toasts."],
        ["Forms and validation", "react-hook-form, @hookform/resolvers and Zod schemas in utils/validation.ts."],
        ["Production backend", "Firebase Auth, Firestore, Storage, Cloud Functions, FCM and Firebase Hosting."],
        ["Server APIs", "Next.js Node runtime route handlers under app/api/* using Firebase Admin and local SQL adapters."],
        ["Local data mode", "SQLite through Node's node:sqlite module in lib/local-sql.ts, selected by DATA_BACKEND=sql."],
        ["Payments", "M-Pesa Daraja STK Push route handlers plus wallet/transaction documents."],
        ["Verification", "Resend email OTP and Didit hosted KYC with signed webhook verification."],
        ["Maps", "Mapbox client picker for address and coordinate capture."],
    ], [1.7, 4.8])
    doc.add_paragraph("Important scripts:")
    add_bullets(doc, ["npm run dev - local development server.", "npm run typecheck - TypeScript validation.", "npm run lint - ESLint.", "npm run build - production Next.js build.", "npm run functions:build - Firebase Functions TypeScript build."])

    doc.add_heading("Mental Model", level=1)
    add_numbered(doc, [
        "A user signs in with Firebase Auth; AuthContext waits for both auth state and profile data.",
        "The role in the user profile determines the home route, navigation and permissions.",
        "Client screens call service modules, and service modules call Next.js API routes or Firebase client SDKs.",
        "API routes verify Firebase ID tokens with Admin SDK, then read/write Firestore or local SQLite depending on data-backend mode.",
        "Important financial and workflow actions use server-side transactions so status, wallet and notification records stay consistent.",
    ])

    doc.add_heading("Source Tree", level=1)
    add_table(doc, ["Folder/file", "Purpose"], [
        ["app/", "Next.js App Router pages, layouts, global CSS, manifest and API route handlers."],
        ["components/", "Reusable UI, auth, job, location, profile, wallet, ratings and PWA components."],
        ["context/AuthContext.tsx", "Global authentication/profile state provider."],
        ["hooks/useProtectedRoute.ts", "Route guards for protected and public-only pages."],
        ["services/", "Client-side functions that hide API calls and Firebase interaction details from pages."],
        ["lib/", "Firebase, admin, SQL, payment, verification, location and server helper libraries."],
        ["utils/", "Pure formatting, pricing, validation, duration, phone and verification helpers."],
        ["types/index.ts", "Central TypeScript domain model definitions."],
        ["functions/", "Firebase Cloud Functions source and build config."],
        ["public/", "PWA service worker, icons, design assets and public images."],
        ["firestore.rules / storage.rules", "Production access-control policy for Firestore and Storage."],
        ["firebase.json / firestore.indexes.json", "Firebase hosting, functions, rules and index deployment configuration."],
    ], [2.0, 4.5])

    doc.add_heading("User-Facing Pages", level=1)
    add_table(doc, ["Route", "Screen", "Purpose"], PAGES, [1.35, 1.55, 3.6])

    doc.add_heading("Main User Workflows", level=1)
    doc.add_heading("Authentication and Role Routing", level=2)
    add_bullets(doc, [
        "AuthContext subscribes to Firebase Auth, then loads the user's profile from Firestore or /api/auth/me in SQL mode.",
        "Recovered role storage prevents temporary blank-profile flicker after auth changes.",
        "useProtectedRoute redirects anonymous users to /auth/login and users without profiles to /complete-profile.",
        "usePublicOnlyRoute blocks signed-in users from login/register pages and sends them to their role home.",
        "Role homes are admin -> /admin, client -> /workers, worker -> /dashboard.",
    ])
    doc.add_heading("Worker Workflow", level=2)
    add_numbered(doc, [
        "Workers browse /jobs and open /jobs/[id] to apply.",
        "Applications are created by POST /api/applications after checking role, lock state, open job status and accepted-worker limits.",
        "Pending applications can be cancelled by the worker; the server tracks daily cancellation counts and blocks after two per day.",
        "Accepted applications unlock chat and appear as live/current work.",
        "Workers mark accepted jobs complete, moving the application to completion_requested and notifying the client.",
        "Direct hire requests are represented as pending rehire applications; workers can accept or reject them.",
    ])
    doc.add_heading("Client Workflow", level=2)
    add_numbered(doc, [
        "Clients post work through /jobs/new or /find-work; job input is validated by utils/validation.ts and server routes.",
        "Clients review applicants on /applications and can accept a worker if wallet escrow is sufficient.",
        "Accepting an application reserves escrow, creates/unlocks a conversation and notifies the worker.",
        "When a worker requests completion, the client confirms in /applications; server logic releases escrow to the worker wallet, records transactions and updates job/application status.",
        "Clients can rehire workers with explicit duration and start-date information.",
    ])
    doc.add_heading("Admin Workflow", level=2)
    add_bullets(doc, [
        "Admin pages are wrapped by app/admin/layout.tsx and protected by admin role/permission checks.",
        "Admins can review KYC, lock/unlock or moderate users, answer support tickets, inspect transactions, reconcile M-Pesa state and view audit logs.",
        "Admin route handlers use lib/admin-security.ts to map roles to permissions and write audit logs for sensitive actions.",
    ])

    doc.add_heading("API Route Map", level=1)
    add_table(doc, ["Methods", "Endpoint", "Responsibility"], APIS, [1.0, 2.1, 3.4])

    doc.add_heading("Data Model", level=1)
    add_table(doc, ["Type", "Meaning"], MODELS, [2.0, 4.5])
    add_note(
        doc,
        "Compatibility note:",
        "Several screens intentionally support both newer detailed records and older fields, for example skillProfiles plus legacy skills[], and rateAmount/paymentType fields beside current payAmount/paymentMethod fields."
    )

    doc.add_heading("Services, Libraries and Utilities", level=1)
    doc.add_heading("Client Services", level=2)
    add_table(doc, ["File", "Responsibility"], SERVICES, [2.25, 4.25])
    doc.add_heading("Shared Libraries", level=2)
    add_table(doc, ["File", "Responsibility"], LIBS, [2.25, 4.25])
    doc.add_heading("Utilities", level=2)
    add_table(doc, ["File", "Responsibility"], UTILS, [2.25, 4.25])

    doc.add_heading("Reusable Components", level=1)
    add_table(doc, ["Component area", "Responsibility"], COMPONENTS, [2.35, 4.15])

    doc.add_heading("Firebase and Security", level=1)
    add_table(doc, ["Security area", "How it works"], [
        ["Firestore users", "Users can create/update only safe profile fields; admin-only fields like role, lock state, wallet counters and ratings are protected."],
        ["Jobs", "Clients manage their own jobs, workers can read open/assigned jobs, and server/API logic handles sensitive status changes."],
        ["Applications", "Workers can create pending applications and withdraw their own pending records; richer state transitions happen server-side."],
        ["Messages", "Only verified participants in unlocked conversations can send/read messages; admins can inspect/delete as needed."],
        ["Wallets/transactions", "Users can read their own wallet/transactions but cannot directly mutate financial records from the browser."],
        ["Support", "Users can create/read their own tickets; admins can read and manage all tickets."],
        ["Backend-only collections", "emailOtps, identityClaims, diditWebhookEvents and mpesaCallbacks are blocked from browser access."],
        ["Storage", "Profile images, verification files, job attachments, skill proofs and message images have owner/admin/participant checks plus file-size/type limits."],
    ], [1.7, 4.8])

    doc.add_heading("Cloud Functions", level=1)
    add_table(doc, ["Function", "Purpose"], [
        ["reviewVerification / reviewKyc", "Admin callable that approves/rejects verification and notifies the user."],
        ["setAdminRole", "Admin callable that sets custom claims and role field."],
        ["acceptApplication", "Legacy callable that accepts an application and creates conversation records."],
        ["markPaidInCash", "Completes cash jobs and locks workers until service fee is paid."],
        ["notifyOnNotificationCreate", "Sends FCM web push when notification records are created."],
        ["alertMatchingWorkersOnJobCreate", "Creates worker notifications/activities for jobs matching skills/category."],
        ["auditUserUpdates", "Logs role, verification or lock-state changes."],
        ["updateBadgesOnJobCompletion", "Updates worker badges after completed jobs."],
        ["requestDeposit/completeMpesaJob/payOutstandingFee", "Legacy callables retained but instruct callers to use current server payment flows."],
        ["mpesaCallback", "Legacy HTTP callback logger; the current app also has /api/mpesa/callback."],
    ], [2.2, 4.3])

    doc.add_heading("Payments and Wallet Accounting", level=1)
    add_bullets(doc, [
        "Deposits begin at /api/mpesa/stk-push, which creates an initiating transaction, calls Daraja, then marks it pending and increments pendingBalance.",
        "Callbacks at /api/mpesa/callback validate provider data, protect against duplicate checkout credits and move funds into available wallet balance.",
        "Accepting a worker reserves escrow by moving client funds from available/balance into pendingBalance and creating a pending payout transaction.",
        "Client completion confirmation releases pending escrow, credits the worker wallet, creates payout/wallet_credit/service_fee records and updates job/application status.",
        "Cash payment flows lock workers for outstanding platform service fees until settlement.",
    ])

    doc.add_heading("Verification and Trust", level=1)
    add_bullets(doc, [
        "Email verification uses Resend-delivered 6-digit OTPs, a 10-minute expiry, resend cooldown, attempt limits and SHA-256 hashes protected by OTP_SECRET.",
        "Identity verification uploads ID front/back and proof-of-address documents to protected Storage paths.",
        "/api/kyc/start creates a Didit hosted session and a pending verification record.",
        "/api/kyc/didit-webhook validates X-Signature-V2 and timestamp freshness before updating identity status.",
        "Admin KYC review finalizes address/document approval or rejection.",
    ])

    doc.add_heading("Local SQL Development Mode", level=1)
    add_bullets(doc, [
        "Set DATA_BACKEND=sql and NEXT_PUBLIC_DATA_BACKEND=sql to use SQLite for app data while keeping Firebase Auth.",
        "lib/local-sql.ts owns schema creation, migrations via ensureColumn, row mappers and local versions of user/job/application/wallet/chat/notification/rating functions.",
        "/api/auth/me can import an existing Firestore profile into local SQL if one exists.",
        "This mode is for local development only; production should use Firebase/Firestore rules and server APIs.",
    ])

    doc.add_heading("Environment Variables", level=1)
    add_table(doc, ["Group", "Variables", "Purpose"], ENV_ROWS, [1.4, 2.3, 2.8])

    doc.add_heading("Deployment", level=1)
    add_bullets(doc, [
        "firebase.json configures framework-aware Firebase Hosting, Firestore rules/indexes, Storage rules and Node.js 20 Functions.",
        "The PWA service worker is served with no-cache headers so updates can be picked up quickly.",
        "Recommended checks before deployment: npm run typecheck, npm run lint, npm run build and npm run functions:build.",
        "Deployment commands: firebase deploy --only firestore:rules,firestore:indexes,storage; npm run functions:build; firebase deploy --only functions,hosting.",
    ])

    doc.add_heading("Important Operational Rules", level=1)
    add_bullets(doc, [
        "Never let browser code mutate wallet balances or transaction status directly.",
        "Keep application status transitions explicit: pending, accepted, completion_requested, completed, rejected, withdrawn.",
        "Preserve legacy data compatibility when changing profiles, skills, jobs or payments.",
        "When auth appears stuck, inspect both Firebase auth state and profile loading state.",
        "When workflow buttons appear missing, inspect both the dedicated /applications page and dashboard modal surfaces.",
        "After edits, run typecheck, lint and build because App Router route boundaries can fail only at build time.",
    ])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Appendix A - File Inventory", level=1)
    doc.add_paragraph("This appendix lists the current project files, excluding generated lock/build/cache artifacts and this documentation output.")
    inventory = file_inventory()
    chunk_size = 35
    for index in range(0, len(inventory), chunk_size):
        if index:
            doc.add_paragraph()
        add_table(doc, ["File", "Role in project"], inventory[index:index + chunk_size], [3.25, 3.25], header_fill=LIGHT_GRAY)

    OUT.unlink(missing_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
